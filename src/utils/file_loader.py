"""Module xử lý và trích xuất nội dung từ các loại file"""

import fitz
import docx
import logging
import os
from typing import BinaryIO, Optional, Union

# Thiết lập logging
logger = logging.getLogger(__name__)

def detect_file_type(file: BinaryIO) -> str:
    """Tự động phát hiện loại file dựa trên định dạng hoặc extension

    Args:
        file: File object đã được upload

    Returns:
        Loại file: "PDF", "Word", hoặc "Unknown"
    """
    filename = getattr(file, 'name', '')
    if filename:
        extension = os.path.splitext(filename)[1].lower()
        if extension == '.pdf':
            return "PDF"
        elif extension in ['.docx', '.doc']:
            return "Word"

    # Nếu không phát hiện được từ tên file, thử đọc nội dung file
    try:
        # Lưu vị trí con trỏ file để đảm bảo có thể đọc lại sau
        current_position = file.tell()

        # Đọc một phần nhỏ đầu file để xác định signature
        file.seek(0)
        header = file.read(8)
        file.seek(current_position)  # Đặt lại vị trí con trỏ

        # Kiểm tra PDF signature
        if header.startswith(b'%PDF'):
            return "PDF"

        # Kiểm tra DOC/DOCX signature (Office Open XML)
        if header.startswith(b'PK\x03\x04'):
            return "Word"

    except Exception as e:
        logger.warning(f"Không thể phát hiện loại file: {str(e)}")

    # Nếu không phát hiện được, thử dựa vào mime type (nếu có)
    mime_type = getattr(file, 'type', '')
    if mime_type:
        if 'pdf' in mime_type.lower():
            return "PDF"
        elif 'word' in mime_type.lower() or 'docx' in mime_type.lower():
            return "Word"

    return "Unknown"

def load_file(file: BinaryIO) -> str:
    """Tải và trích xuất nội dung từ file

    Args:
        file: File object được upload từ Streamlit

    Returns:
        Nội dung văn bản từ file hoặc thông báo lỗi
    """
    try:
        # Phát hiện loại file
        file_type = detect_file_type(file)
        logger.info(f"Đã phát hiện loại file: {file_type}")

        if file_type == "PDF":
            return extract_text_from_pdf(file)
        elif file_type == "Word":
            return extract_text_from_word(file)
        else:
            return f"Lỗi: Không hỗ trợ định dạng file này. Vui lòng tải lên file PDF hoặc Word (.docx)."

    except Exception as e:
        logger.error(f"Lỗi khi đọc file: {str(e)}")
        return f"Lỗi khi xử lý file: {str(e)}"

def extract_text_from_pdf(file: BinaryIO) -> str:
    """Trích xuất văn bản từ file PDF

    Args:
        file: File PDF

    Returns:
        Nội dung văn bản đã trích xuất
    """
    text = ""
    try:
        # Lưu vị trí hiện tại
        current_position = file.tell()
        file.seek(0)  # Đặt lại vị trí đầu file

        # Mở và xử lý file PDF
        with fitz.open(stream=file.read(), filetype="pdf") as doc:
            # Số trang trong file
            page_count = len(doc)
            logger.info(f"Đang xử lý file PDF: {page_count} trang")

            # Trích xuất văn bản từ mỗi trang
            for i, page in enumerate(doc):
                page_text = page.get_text("text")
                text += page_text + "\n\n"

                # Log tiến trình
                if i % 10 == 0 and i > 0:
                    logger.info(f"Đã xử lý {i}/{page_count} trang")

        # Đặt lại vị trí con trỏ file
        file.seek(current_position)

        return text.strip()

    except Exception as e:
        logger.error(f"Lỗi khi trích xuất văn bản từ PDF: {str(e)}")
        raise

def extract_text_from_word(file: BinaryIO) -> str:
    """Trích xuất văn bản từ file Word (.docx)

    Args:
        file: File Word

    Returns:
        Nội dung văn bản đã trích xuất
    """
    try:
        # Đọc tài liệu Word
        doc = docx.Document(file)

        # Trích xuất văn bản từ các đoạn
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]

        # Trích xuất văn bản từ các bảng
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables_text.append(row_text)

        # Kết hợp tất cả văn bản
        all_text = "\n\n".join(paragraphs)
        if tables_text:
            all_text += "\n\n" + "\n".join(tables_text)

        return all_text.strip()

    except Exception as e:
        logger.error(f"Lỗi khi trích xuất văn bản từ Word: {str(e)}")
        raise